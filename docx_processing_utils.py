import re
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE
from text_processing_utils import replace_with_placeholders, needs_translation
from config import MATH_LIKE_CHARS

def is_equation_run(run):
    return bool(run._element.xpath('.//m:oMath') or run._element.xpath('.//m:oMathPara'))

def has_special_formatting(run):
    try:
        if run.font.subscript or run.font.superscript or run.font.bold or run.font.italic:
            return True
    except AttributeError:
        pass
    return False

def split_text_by_proportions(text_to_split, original_runs, original_full_text, placeholder_map_for_segment):
    if not original_runs or not text_to_split.strip():
        return [""] * len(original_runs)
    is_placeholder_for_full_segment = text_to_split in placeholder_map_for_segment and placeholder_map_for_segment[text_to_split] == original_full_text
    if len(original_runs) == 1 or is_placeholder_for_full_segment:
        return [text_to_split] + [""] * (len(original_runs) - 1)
    original_run_texts = [r.text for r in original_runs]
    total_original_char_length = sum(len(rt) for rt in original_run_texts)
    if total_original_char_length == 0:
        num_runs = len(original_runs)
        avg_len = len(text_to_split) // num_runs if num_runs > 0 else len(text_to_split)
        parts = []
        current_pos = 0
        for i in range(num_runs):
            if i == num_runs -1:
                parts.append(text_to_split[current_pos:])
            else:
                parts.append(text_to_split[current_pos : current_pos + avg_len])
            current_pos += avg_len
        return parts
    split_parts = []
    remaining_text_to_split = text_to_split
    for i, run_text in enumerate(original_run_texts):
        if not remaining_text_to_split:
            split_parts.append("")
            continue
        if i == len(original_runs) - 1:
            split_parts.append(remaining_text_to_split)
            remaining_text_to_split = ""
            continue
        proportion = len(run_text) / total_original_char_length if total_original_char_length > 0 else (1/len(original_runs))
        estimated_len = max(0, int(len(remaining_text_to_split) * proportion)) if proportion > 0 else 0
        search_radius = min(10, int(len(remaining_text_to_split) * 0.2))
        split_at = -1
        for j in range(estimated_len + search_radius, estimated_len - search_radius -1, -1):
            if 0 < j < len(remaining_text_to_split) and remaining_text_to_split[j] == ' ':
                split_at = j
                break
        if split_at == -1:
            if estimated_len == 0 and len(run_text) > 0 :
                split_at = 1 if remaining_text_to_split else 0
            elif estimated_len >= len(remaining_text_to_split):
                split_at = len(remaining_text_to_split)
            else:
                split_at = estimated_len
        split_at = max(0, min(split_at, len(remaining_text_to_split)))
        split_parts.append(remaining_text_to_split[:split_at])
        remaining_text_to_split = remaining_text_to_split[split_at:]
    if remaining_text_to_split and split_parts:
        split_parts[-1] += remaining_text_to_split
    elif remaining_text_to_split and not split_parts:
        return [remaining_text_to_split] + [""] * (len(original_runs) -1)
    while len(split_parts) < len(original_runs):
        split_parts.append("")
    return split_parts[:len(original_runs)]

def _get_run_formatting(run):
    fmt = {
        'bold': None, 'italic': None, 'underline': None, 'strike': None,
        'superscript': None, 'subscript': None, 'shadow': None, 'outline': None,
        'imprint': None, 'emboss': None, 'size': None, 'name': None, 'color_rgb': None,
        'all_caps': None, 'small_caps': None
    }
    try:
        font = run.font
        fmt['bold'] = font.bold
        fmt['italic'] = font.italic
        fmt['underline'] = font.underline
        fmt['strike'] = font.strike
        fmt['superscript'] = font.superscript
        fmt['subscript'] = font.subscript
        if font.size is not None:
            fmt['size'] = font.size
        fmt['name'] = font.name
        if font.color and hasattr(font.color, 'rgb') and font.color.rgb is not None:
            fmt['color_rgb'] = font.color.rgb
    except AttributeError:
        pass
    return fmt

def collect_texts_and_images(doc, entities):
    texts_collected = []
    
    def process_paragraph_runs(para, para_type_prefix, unique_para_id):
        current_standard_group = []
        current_math_group = []
        for run_idx_in_para, run in enumerate(para.runs):
            if is_equation_run(run):
                if current_standard_group:
                    group_text = " ".join(r.text for r in current_standard_group).strip()
                    if group_text:
                        mod_text, p_map, mws = replace_with_placeholders(group_text, entities)
                        fmt_info = [_get_run_formatting(r) for r in current_standard_group]
                        texts_collected.append(("body", f"{para_type_prefix}_{unique_para_id}_std_{run_idx_in_para-len(current_standard_group)}", para, group_text, mod_text, p_map, needs_translation(mod_text), mws, list(current_standard_group), fmt_info))
                    current_standard_group = []
                if current_math_group:
                    group_text = "".join(r.text for r in current_math_group)
                    if group_text.strip():
                        mod_text, p_map, mws = replace_with_placeholders(group_text, entities)
                        fmt_info = [_get_run_formatting(r) for r in current_math_group]
                        texts_collected.append(("math_text", f"{para_type_prefix}_{unique_para_id}_math_{run_idx_in_para-len(current_math_group)}", para, group_text, mod_text, p_map, needs_translation(mod_text), mws, list(current_math_group), fmt_info))
                    current_math_group = []
                texts_collected.append(("omath_object", f"{para_type_prefix}_{unique_para_id}_eq_{run_idx_in_para}", para, run.text, run.text, {}, False, [], [run], [_get_run_formatting(run)]))
                continue

            if run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict'):
                if current_standard_group:
                    group_text = " ".join(r.text for r in current_standard_group).strip()
                    if group_text:
                        mod_text, p_map, mws = replace_with_placeholders(group_text, entities)
                        fmt_info = [_get_run_formatting(r) for r in current_standard_group]
                        texts_collected.append(("body", f"{para_type_prefix}_{unique_para_id}_std_{run_idx_in_para-len(current_standard_group)}", para, group_text, mod_text, p_map, needs_translation(mod_text), mws, list(current_standard_group), fmt_info))
                    current_standard_group = []
                if current_math_group:
                    group_text = "".join(r.text for r in current_math_group)
                    if group_text.strip():
                        mod_text, p_map, mws = replace_with_placeholders(group_text, entities)
                        fmt_info = [_get_run_formatting(r) for r in current_math_group]
                        texts_collected.append(("math_text", f"{para_type_prefix}_{unique_para_id}_math_{run_idx_in_para-len(current_math_group)}", para, group_text, mod_text, p_map, needs_translation(mod_text), mws, list(current_math_group), fmt_info))
                    current_math_group = []
                texts_collected.append(("preserved_drawing", f"{para_type_prefix}_{unique_para_id}_draw_{run_idx_in_para}", para, run.text, run.text, {}, False, [], [run], [_get_run_formatting(run)]))
                continue
            
            run_text_stripped = run.text.strip()
            is_mathy_char_run = run_text_stripped in MATH_LIKE_CHARS and len(run_text_stripped) == 1
            if has_special_formatting(run) or is_mathy_char_run:
                if current_standard_group:
                    group_text = " ".join(r.text for r in current_standard_group).strip()
                    if group_text:
                        mod_text, p_map, mws = replace_with_placeholders(group_text, entities)
                        fmt_info = [_get_run_formatting(r) for r in current_standard_group]
                        texts_collected.append(("body", f"{para_type_prefix}_{unique_para_id}_std_{run_idx_in_para-len(current_standard_group)}", para, group_text, mod_text, p_map, needs_translation(mod_text), mws, list(current_standard_group), fmt_info))
                    current_standard_group = []
                current_math_group.append(run)
            else:
                if current_math_group:
                    group_text = "".join(r.text for r in current_math_group)
                    if group_text.strip():
                        mod_text, p_map, mws = replace_with_placeholders(group_text, entities)
                        fmt_info = [_get_run_formatting(r) for r in current_math_group]
                        texts_collected.append(("math_text", f"{para_type_prefix}_{unique_para_id}_math_{run_idx_in_para-len(current_math_group)}", para, group_text, mod_text, p_map, needs_translation(mod_text), mws, list(current_math_group), fmt_info))
                    current_math_group = []
                current_standard_group.append(run)

        if current_standard_group:
            group_text = " ".join(r.text for r in current_standard_group).strip()
            if group_text:
                mod_text, p_map, mws = replace_with_placeholders(group_text, entities)
                fmt_info = [_get_run_formatting(r) for r in current_standard_group]
                texts_collected.append(("body", f"{para_type_prefix}_{unique_para_id}_std_end", para, group_text, mod_text, p_map, needs_translation(mod_text), mws, list(current_standard_group), fmt_info))
        if current_math_group:
            group_text = "".join(r.text for r in current_math_group)
            if group_text.strip():
                mod_text, p_map, mws = replace_with_placeholders(group_text, entities)
                fmt_info = [_get_run_formatting(r) for r in current_math_group]
                texts_collected.append(("math_text", f"{para_type_prefix}_{unique_para_id}_math_end", para, group_text, mod_text, p_map, needs_translation(mod_text), mws, list(current_math_group), fmt_info))

    for para_idx, para in enumerate(doc.paragraphs):
        process_paragraph_runs(para, "body_para", para_idx)
    for section_idx, section in enumerate(doc.sections):
        for header_idx, header_para in enumerate(section.header.paragraphs):
            process_paragraph_runs(header_para, f"s{section_idx}_hdr", header_idx)
        for footer_idx, footer_para in enumerate(section.footer.paragraphs):
            process_paragraph_runs(footer_para, f"s{section_idx}_ftr", footer_idx)
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx_in_cell, para_in_cell in enumerate(cell.paragraphs):
                    process_paragraph_runs(para_in_cell, f"t{table_idx}_r{row_idx}_c{cell_idx}_p", para_idx_in_cell)
    return texts_collected

def apply_formatting(target_runs, text_parts, formatting_info_list):
    for i, run in enumerate(target_runs):
        text_content_for_run = text_parts[i] if i < len(text_parts) else ""
        run.text = text_content_for_run
        if i < len(formatting_info_list):
            fmt = formatting_info_list[i]
            font = run.font
            try:
                if fmt.get('bold') is not None: font.bold = fmt['bold']
                if fmt.get('italic') is not None: font.italic = fmt['italic']
                underline_val = fmt.get('underline')
                if underline_val is not None:
                    if isinstance(underline_val, bool): font.underline = underline_val
                    else: font.underline = underline_val 
                else: font.underline = False 
                if fmt.get('strike') is not None: font.strike = fmt['strike']
                if fmt.get('superscript') is not None: font.superscript = fmt['superscript']
                if fmt.get('subscript') is not None: font.subscript = fmt['subscript']
                if fmt.get('name') is not None: font.name = fmt['name']
                if fmt.get('size') is not None and fmt['size'] is not None:
                    current_size = fmt['size']
                    if isinstance(current_size, Pt): font.size = current_size
                    elif isinstance(current_size, (int, float)) and current_size > 0 : font.size = Pt(current_size)
                if fmt.get('color_rgb') is not None and isinstance(fmt['color_rgb'], RGBColor):
                    try: font.color.rgb = fmt['color_rgb']
                    except Exception: pass
                elif fmt.get('color_rgb') is None and hasattr(font.color, 'rgb') and font.color.rgb is not None:
                    pass
            except (AttributeError, ValueError, Exception): pass